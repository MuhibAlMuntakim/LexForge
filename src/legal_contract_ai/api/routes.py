from fastapi import APIRouter, UploadFile, File, HTTPException, BackgroundTasks, Response, Depends
import os
import uuid
import logging
import datetime
import re
import io
import json
from difflib import SequenceMatcher
from typing import List, Optional, TYPE_CHECKING
from src.legal_contract_ai.api.schemas import (
    ReviewResponse,
    ReviewRequest,
    ContractHistoryItem,
    ContractSummaryItem,
    DeletedContractItem,
    ContractDeleteResponse,
    ContractRestoreResponse,
    ChatQueryRequest,
    ChatQueryResponse,
    ChatIndexResponse,
    DashboardMetricsResponse,
    AnalyticsMetricsResponse,
    NegotiationItem,
    NegotiationStatusUpdateRequest,
    NegotiationStatusUpdateResponse,
    RedlineAcceptRequest,
    RedlineAcceptResponse,
    RedlineSuggestionUpdateRequest,
    RedlineSuggestionUpdateResponse,
    PlaybookCreateRequest,
    PlaybookUpdateRequest,
    PlaybookResponse,
)
from src.legal_contract_ai.core.contract_pipeline import ContractPipeline
from src.legal_contract_ai.core.contract_memory import ContractMemory
from src.legal_contract_ai.core.metrics_service import MetricsService
from src.legal_contract_ai.core.playbook_service import PlaybookService
from src.legal_contract_ai.analysis.risk_scorer import RiskScorer
from src.legal_contract_ai.core.violation_utils import deduplicate_violations
from src.legal_contract_ai.api.security import (
    SecurityContext,
    get_security_context,
    rate_limit_chat,
    rate_limit_review,
)

if TYPE_CHECKING:
    from src.legal_contract_ai.rag.rag_service import RAGService

router = APIRouter()
pipeline = ContractPipeline()
memory = ContractMemory()
ragsvc: Optional["RAGService"] = None
ragsvc_error: Optional[str] = None
metrics_service = MetricsService()
playbook_service = PlaybookService()
logger = logging.getLogger(__name__)

UPLOAD_DIR = "data/uploaded_contracts"
os.makedirs(UPLOAD_DIR, exist_ok=True)

ALLOWED_UPLOAD_EXTENSIONS = {".pdf", ".docx", ".txt"}
ALLOWED_UPLOAD_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}
MAX_UPLOAD_BYTES = 10 * 1024 * 1024


def _save_validated_upload(file: UploadFile, prefix_with_uuid: bool = False) -> tuple[str, str]:
    original_name = os.path.basename(file.filename or "upload.bin")
    ext = os.path.splitext(original_name)[1].lower()
    if ext not in ALLOWED_UPLOAD_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    content_type = str(file.content_type or "").lower()
    if content_type and content_type not in ALLOWED_UPLOAD_CONTENT_TYPES:
        raise HTTPException(status_code=400, detail="Unsupported content type")

    payload = file.file.read(MAX_UPLOAD_BYTES + 1)
    if len(payload) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large")

    safe_name = f"{uuid.uuid4()}_{original_name}" if prefix_with_uuid else f"{uuid.uuid4()}{ext}"
    upload_path = os.path.join(UPLOAD_DIR, safe_name)
    with open(upload_path, "wb") as buffer:
        buffer.write(payload)

    return upload_path, original_name


def _get_rag_service() -> Optional["RAGService"]:
    global ragsvc, ragsvc_error
    if ragsvc is not None:
        return ragsvc
    if ragsvc_error:
        return None

    try:
        from src.legal_contract_ai.rag.rag_service import RAGService

        ragsvc = RAGService()
        return ragsvc
    except Exception as e:
        ragsvc_error = str(e)
        logger.warning(f"RAG service unavailable: {e}")
        return None


def _fallback_suggestion_for_clause(clause: str, policy: str = "") -> str:
    clause_key = (clause or "").strip().lower()
    defaults = {
        "indemnity": (
            "MUTUAL INDEMNIFICATION. Each party shall indemnify, defend, and hold harmless the other party "
            "from third-party claims to the extent arising from the indemnifying party's negligence, willful "
            "misconduct, or breach of this Agreement."
        ),
        "termination": (
            "TERMINATION FOR CONVENIENCE. Customer may terminate on thirty (30) days' written notice without "
            "early termination fees. Vendor may terminate only for uncured material breach after notice and cure period."
        ),
        "liability": (
            "LIMITATION OF LIABILITY. Each party's aggregate liability shall not exceed 100% of annual contract value, "
            "and neither party shall be liable for indirect, special, incidental, or consequential damages."
        ),
        "confidentiality": (
            "CONFIDENTIALITY. Confidentiality obligations apply mutually and survive termination for at least three (3) years."
        ),
    }
    if clause_key in defaults:
        return defaults[clause_key]
    if policy:
        return f"Replace this clause with language aligned to policy: {policy}"
    return "Replace this clause with language aligned to the selected playbook requirements."


def _is_placeholder_suggestion(text: str) -> bool:
    normalized = str(text or "").strip().lower()
    if not normalized:
        return True
    return normalized in {
        "null",
        "none",
        "n/a",
        "na",
        "nil",
        "undefined",
        "not available",
        "not applicable",
    }


def _sanitize_redline_suggestion(text: str) -> str:
    cleaned = str(text or "").strip()
    if not cleaned:
        return ""

    cleaned = re.sub(r"^```[a-zA-Z]*\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    cleaned = re.sub(r"\n\s*(source|sources|citation|citations)\s*:[\s\S]*$", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s+(source|sources|citation|citations)\s*:\s*.*$", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*(\(|\[)?\s*chunks?\s*\d+(?:\s*[-,]\s*\d+)*(\)|\])?\s*$", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*\((?:chunk|chunks)\s*\d+(?:\s*[-,]\s*\d+)*\)\s*$", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(
        r"\n+\s*[\(\[]?\s*(?:redline|chunk|remaining\s+violations?|executive\s+summary)\b[\s\S]*$",
        "",
        cleaned,
        flags=re.IGNORECASE,
    )
    cleaned = re.sub(
        r"\s+[\(\[]\s*(?:redline|chunk|remaining\s+violations?|executive\s+summary)\b[\s\S]*$",
        "",
        cleaned,
        flags=re.IGNORECASE,
    )

    return cleaned.strip().strip('"').strip("'").strip()


def _ensure_contract_redlines_have_suggestions(analysis: dict) -> dict:
    redlines = analysis.get("redlines") or []
    violations = deduplicate_violations(analysis.get("violations") or [])
    if len(violations) != len(analysis.get("violations") or []):
        analysis["violations"] = violations
        analysis["risk_assessment"] = RiskScorer().calculate_score(violations)
    if not redlines:
        return analysis

    # Map first policy/template by clause for fallback enrichment.
    by_clause: dict[str, dict] = {}
    for violation in violations:
        clause = str(violation.get("clause") or "").strip().lower()
        if not clause or clause in by_clause:
            continue
        by_clause[clause] = {
            "policy": str(violation.get("policy") or "").strip(),
            "template": str(violation.get("playbook_template") or "").strip(),
        }

    changed = False
    for item in redlines:
        clause = str(item.get("violation") or "").strip().lower()
        raw_suggested = str(item.get("suggested_redline") or item.get("suggested_clause") or "")
        suggested = _sanitize_redline_suggestion(raw_suggested)
        if suggested and not _is_placeholder_suggestion(suggested):
            if suggested != raw_suggested.strip():
                item["suggested_redline"] = suggested
                changed = True
            continue
        fallback = by_clause.get(clause, {})
        item["suggested_redline"] = (
            fallback.get("template")
            or _fallback_suggestion_for_clause(clause, fallback.get("policy", ""))
        )
        if not item.get("rationale"):
            policy_text = fallback.get("policy", "")
            item["rationale"] = (
                f"Playbook requirement: {policy_text}" if policy_text else "Playbook policy alignment required."
            )
        changed = True

    if changed:
        analysis["redlines"] = redlines
    return analysis

@router.post("/contracts/analyze", response_model=ReviewResponse)
async def analyze_contract(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    playbook_id: str = "default",
    ctx: SecurityContext = Depends(get_security_context),
    _limit: None = Depends(rate_limit_review),
):
    company_id = ctx.company_id
    uploaded_path, original_filename = _save_validated_upload(file)
    
    try:
        try:
            resolved = playbook_service.resolve_for_review(company_id=company_id, playbook_id=playbook_id)
        except FileNotFoundError:
            raise HTTPException(status_code=404, detail="Playbook not found")

        result = pipeline.run_review(
            uploaded_path,
            company_id=company_id,
            playbook_path=resolved["path"],
            playbook_id=resolved["playbook_id"],
            playbook_name=resolved["name"],
        )

        # Index AI review artifacts in background so review response is not blocked.
        background_tasks.add_task(
            _safe_index_analysis_artifacts,
            company_id=company_id,
            contract_id=result["contract_id"],
            analysis=result,
            filename=original_filename,
        )
        
        return result
    except Exception as e:
        logger.exception("Analysis failed")
        raise HTTPException(status_code=500, detail="Internal server error")

# Backward compatibility alias
@router.post("/review", response_model=ReviewResponse)
async def review_contract_alias(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    playbook_id: str = "default",
    ctx: SecurityContext = Depends(get_security_context),
    _limit: None = Depends(rate_limit_review),
):
    return await analyze_contract(background_tasks, file, playbook_id, ctx, _limit)


@router.get("/playbooks", response_model=List[PlaybookResponse])
async def list_playbooks(ctx: SecurityContext = Depends(get_security_context)):
    try:
        return playbook_service.list_playbooks(ctx.company_id)
    except Exception as e:
        logger.exception("List playbooks failed")
        raise HTTPException(status_code=500, detail="Internal server error")


@router.get("/playbooks/{playbook_id}", response_model=PlaybookResponse)
async def get_playbook(playbook_id: str, ctx: SecurityContext = Depends(get_security_context)):
    try:
        return playbook_service.get_playbook(ctx.company_id, playbook_id)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Playbook not found")
    except Exception as e:
        logger.exception("Get playbook failed")
        raise HTTPException(status_code=500, detail="Internal server error")


@router.post("/playbooks", response_model=PlaybookResponse)
async def create_playbook(payload: PlaybookCreateRequest, ctx: SecurityContext = Depends(get_security_context)):
    try:
        return playbook_service.create_playbook(
            company_id=ctx.company_id,
            name=payload.name,
            description=payload.description or "",
            rules=payload.rules,
            playbook_version=payload.playbook_version,
        )
    except Exception as e:
        logger.exception("Create playbook failed")
        raise HTTPException(status_code=500, detail="Internal server error")


@router.put("/playbooks/{playbook_id}", response_model=PlaybookResponse)
async def update_playbook(
    playbook_id: str,
    payload: PlaybookUpdateRequest,
    ctx: SecurityContext = Depends(get_security_context),
):
    try:
        return playbook_service.update_playbook(
            company_id=ctx.company_id,
            playbook_id=playbook_id,
            name=payload.name,
            description=payload.description,
            rules=payload.rules,
            playbook_version=payload.playbook_version,
        )
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Playbook not found")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.exception("Update playbook failed")
        raise HTTPException(status_code=500, detail="Internal server error")


@router.delete("/playbooks/{playbook_id}")
async def delete_playbook(playbook_id: str, ctx: SecurityContext = Depends(get_security_context)):
    try:
        playbook_service.delete_playbook(company_id=ctx.company_id, playbook_id=playbook_id)
        return {"deleted": True, "playbook_id": playbook_id}
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Playbook not found")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.exception("Delete playbook failed")
        raise HTTPException(status_code=500, detail="Internal server error")


def _safe_index_analysis_artifacts(
    company_id: str,
    contract_id: str,
    analysis: dict,
    filename: str,
) -> None:
    rag_service = _get_rag_service()
    if rag_service is None:
        return

    try:
        rag_service.index_analysis_artifacts(
            company_id=company_id,
            contract_id=contract_id,
            analysis=analysis,
            filename=filename,
        )
    except Exception as rag_error:
        logger.warning(f"Failed to index analysis artifacts for RAG: {rag_error}")


def _build_analysis_context_for_chat(analysis: dict) -> str:
    summary = str(analysis.get("summary") or "")
    document_text = str(analysis.get("document_text") or "")
    redlines = analysis.get("redlines") or []
    violations = analysis.get("violations") or []
    risk = analysis.get("risk_assessment") or {}
    accepted_indexes = set(analysis.get("accepted_redline_indexes") or [])
    rejected_indexes = set(analysis.get("rejected_redline_indexes") or [])

    resolved_clauses: set[str] = set()
    for idx in accepted_indexes:
        if idx < 0 or idx >= len(redlines):
            continue
        clause = str(redlines[idx].get("violation") or "").strip().lower()
        if clause:
            resolved_clauses.add(clause)

    remaining_violations = []
    for item in violations:
        clause = str(item.get("clause") or "").strip().lower()
        if clause and clause in resolved_clauses:
            continue
        remaining_violations.append(item)

    redline_lines = []
    for i, item in enumerate(redlines[:20], start=1):
        original = str(item.get("original_clause") or "").strip()
        suggested = str(item.get("suggested_redline") or item.get("suggested_clause") or "").strip()
        rationale = str(item.get("rationale") or "").strip()
        redline_lines.append(
            f"Redline {i}:\\nOriginal: {original}\\nSuggested: {suggested}\\nRationale: {rationale}"
        )

    violation_lines = []
    for item in remaining_violations[:20]:
        clause = str(item.get("clause") or "Unknown")
        sev = str(item.get("severity") or "medium")
        issue = str(item.get("issue") or "")
        violation_lines.append(f"[{sev}] {clause}: {issue}")

    doc_excerpt = document_text[:12000]
    return "\\n\\n".join(
        [
            "SAVED CONTRACT ANALYSIS CONTEXT",
            "LIVE REVIEW STATE (AUTHORITATIVE FOR CURRENT STATUS)",
            f"Accepted suggestions: {len(accepted_indexes)}",
            f"Rejected suggestions: {len(rejected_indexes)}",
            f"Remaining violations count: {len(remaining_violations)}",
            f"Summary: {summary}",
            f"Risk: {json.dumps(risk)}",
            "Remaining Violations:\\n" + ("\\n".join(violation_lines) if violation_lines else "None"),
            "Redlines:\\n" + ("\\n\\n".join(redline_lines) if redline_lines else "None"),
            "Document Excerpt:\\n" + doc_excerpt,
        ]
    )


def _answer_from_saved_analysis(
    rag_service: "RAGService",
    payload: ChatQueryRequest,
) -> Optional[dict]:
    contract_id = (payload.contract_id or "").strip()
    if not contract_id:
        return None

    analysis = memory.get_analysis(contract_id)
    if not analysis:
        return None

    context = _build_analysis_context_for_chat(analysis)
    history = payload.chat_history or []
    history_str = "\\n".join(f"{m.role}: {m.content}" for m in history[-8:])

    chain = rag_service.prompt | rag_service.llm
    response = chain.invoke(
        {
            "question": payload.question,
            "chat_history": history_str or "(none)",
            "context": context,
        }
    )

    return {
        "question": payload.question,
        "answer": str(response.content or "").strip(),
        "citations": [
            {
                "chunk_id": f"analysis:{contract_id}",
                "doc_id": None,
                "filename": "saved_contract_analysis",
                "contract_id": contract_id,
                "score": None,
                "excerpt": context[:400],
            }
        ],
        "grounded": False,
    }


def _build_company_context_for_chat(company_id: str) -> Optional[str]:
    analyses = metrics_service.load_analyses()
    company_analyses = [a for a in analyses if str(a.get("company_id", "")).strip() == str(company_id or "").strip()]
    if not company_analyses:
        return None

    total_contracts = len(company_analyses)
    negotiation_count = 0
    reviewed_count = 0
    pending_manual_count = 0
    approved_count = 0
    high_risk_count = 0

    total_risk = 0
    violation_totals = {"high": 0, "medium": 0, "low": 0}

    recent_lines = []
    for item in company_analyses[:12]:
        cid = str(item.get("contract_id", ""))
        short_id = cid[:8] if cid else "unknown"

        risk = item.get("risk_assessment", {}) or {}
        score = int(risk.get("risk_score", 0) or 0)
        level = str(risk.get("risk_level", "medium")).lower()
        counts = risk.get("violation_counts", {}) or {}
        rem_high = int(counts.get("high", 0) or 0)
        rem_medium = int(counts.get("medium", 0) or 0)
        rem_low = int(counts.get("low", 0) or 0)
        rem_total = rem_high + rem_medium + rem_low

        total_risk += score
        violation_totals["high"] += rem_high
        violation_totals["medium"] += rem_medium
        violation_totals["low"] += rem_low

        status = metrics_service._derive_status(item)
        if status == "negotiating":
            negotiation_count += 1
        elif status == "approved":
            approved_count += 1
        elif status in {"reviewed", "in_review"}:
            reviewed_count += 1

        if item.get("requires_manual_review"):
            pending_manual_count += 1
        if score > 60 or level == "high":
            high_risk_count += 1

        recent_lines.append(
            f"- Contract {short_id}: status={status}, risk_score={score}, residual_violations={rem_total}, high={rem_high}, medium={rem_medium}, low={rem_low}"
        )

    avg_risk = round(total_risk / max(total_contracts, 1), 1)

    return "\n\n".join(
        [
            "COMPANY CONTRACT PORTFOLIO CONTEXT",
            f"company_id: {company_id}",
            f"total_contracts: {total_contracts}",
            f"negotiating_contracts: {negotiation_count}",
            f"reviewed_contracts: {reviewed_count}",
            f"approved_contracts: {approved_count}",
            f"requires_manual_review_contracts: {pending_manual_count}",
            f"high_risk_contracts: {high_risk_count}",
            f"average_risk_score: {avg_risk}",
            (
                "portfolio_residual_violations: "
                f"high={violation_totals['high']}, medium={violation_totals['medium']}, low={violation_totals['low']}"
            ),
            "RECENT CONTRACT SNAPSHOT:\n" + "\n".join(recent_lines),
        ]
    )


def _build_company_metrics_snapshot(company_id: str) -> Optional[dict]:
    analyses = metrics_service.load_analyses()
    company_analyses = [
        a
        for a in analyses
        if str(a.get("company_id", "")).strip() == str(company_id or "").strip()
    ]
    if not company_analyses:
        return None

    snapshot = {
        "company_id": company_id,
        "total_contracts": len(company_analyses),
        "negotiating_contracts": 0,
        "reviewed_contracts": 0,
        "approved_contracts": 0,
        "requires_manual_review_contracts": 0,
        "high_risk_contracts": 0,
        "average_risk_score": 0.0,
        "portfolio_residual_violations": {"high": 0, "medium": 0, "low": 0},
    }

    total_risk = 0
    for item in company_analyses:
        risk = item.get("risk_assessment", {}) or {}
        score = int(risk.get("risk_score", 0) or 0)
        level = str(risk.get("risk_level", "medium")).lower()
        counts = risk.get("violation_counts", {}) or {}

        snapshot["portfolio_residual_violations"]["high"] += int(counts.get("high", 0) or 0)
        snapshot["portfolio_residual_violations"]["medium"] += int(counts.get("medium", 0) or 0)
        snapshot["portfolio_residual_violations"]["low"] += int(counts.get("low", 0) or 0)

        status = metrics_service._derive_status(item)
        if status == "negotiating":
            snapshot["negotiating_contracts"] += 1
        elif status == "approved":
            snapshot["approved_contracts"] += 1
        elif status in {"reviewed", "in_review"}:
            snapshot["reviewed_contracts"] += 1

        if item.get("requires_manual_review"):
            snapshot["requires_manual_review_contracts"] += 1
        if score > 60 or level == "high":
            snapshot["high_risk_contracts"] += 1

        total_risk += score

    snapshot["average_risk_score"] = round(total_risk / max(snapshot["total_contracts"], 1), 1)
    return snapshot


def _try_answer_company_metrics_question(payload: ChatQueryRequest) -> Optional[dict]:
    if payload.contract_id:
        return None

    question = str(payload.question or "").strip().lower()
    if not question:
        return None

    snapshot = _build_company_metrics_snapshot(payload.company_id)
    if not snapshot:
        return None

    def response(answer: str) -> dict:
        return {
            "question": payload.question,
            "answer": answer,
            "citations": [
                {
                    "chunk_id": f"company-metrics:{payload.company_id}",
                    "doc_id": None,
                    "filename": "company_metrics_snapshot",
                    "contract_id": None,
                    "score": None,
                    "excerpt": json.dumps(snapshot)[:400],
                }
            ],
            "grounded": False,
        }

    asks_count = any(token in question for token in ["how many", "count", "number of", "total"])
    if asks_count and any(token in question for token in ["negotiating", "negotiation"]):
        return response(
            f"Negotiating contracts for {snapshot['company_id']}: {snapshot['negotiating_contracts']}"
        )
    if asks_count and any(token in question for token in ["reviewed", "in review"]):
        return response(
            f"Reviewed contracts for {snapshot['company_id']}: {snapshot['reviewed_contracts']}"
        )
    if asks_count and any(token in question for token in ["approved", "accepted"]):
        return response(
            f"Approved contracts for {snapshot['company_id']}: {snapshot['approved_contracts']}"
        )
    if asks_count and any(token in question for token in ["manual review", "pending"]):
        return response(
            f"Contracts requiring manual review for {snapshot['company_id']}: {snapshot['requires_manual_review_contracts']}"
        )
    if asks_count and "high risk" in question:
        return response(
            f"High-risk contracts for {snapshot['company_id']}: {snapshot['high_risk_contracts']}"
        )
    if asks_count and "contract" in question:
        return response(
            f"Total contracts for {snapshot['company_id']}: {snapshot['total_contracts']}"
        )

    if any(token in question for token in ["metrics", "summary", "overview", "dashboard"]):
        v = snapshot["portfolio_residual_violations"]
        return response(
            " | ".join(
                [
                    f"Company: {snapshot['company_id']}",
                    f"Total: {snapshot['total_contracts']}",
                    f"Negotiating: {snapshot['negotiating_contracts']}",
                    f"Reviewed: {snapshot['reviewed_contracts']}",
                    f"Approved: {snapshot['approved_contracts']}",
                    f"Manual Review: {snapshot['requires_manual_review_contracts']}",
                    f"High Risk: {snapshot['high_risk_contracts']}",
                    f"Avg Risk Score: {snapshot['average_risk_score']}",
                    f"Residual Violations (H/M/L): {v['high']}/{v['medium']}/{v['low']}",
                ]
            )
        )

    return None

@router.get("/history", response_model=List[ContractHistoryItem])
async def get_history():
    if not os.path.exists("data/contracts"):
        return []
    files = os.listdir("data/contracts")
    history = []
    for f in files:
        if f.endswith(".json"):
            data = memory.get_analysis(f.replace(".json", ""))
            if data:
                history.append({
                    "contract_id": data["contract_id"],
                    "company_id": data["company_id"],
                    "timestamp": data["timestamp"]
                })
    return history


@router.get("/contracts/summaries", response_model=List[ContractSummaryItem])
async def get_contract_summaries(timeframe: str = "all"):
    analyses = metrics_service._filter_analyses_by_timeframe(metrics_service.load_analyses(), timeframe)

    summaries: List[ContractSummaryItem] = []
    for data in analyses:
        risk_assessment = data.get("risk_assessment") or {}
        violation_counts = risk_assessment.get("violation_counts") or {}

        residual_violations = (
            int(violation_counts.get("high", 0) or 0)
            + int(violation_counts.get("medium", 0) or 0)
            + int(violation_counts.get("low", 0) or 0)
        )
        if residual_violations <= 0:
            residual_violations = len(data.get("violations") or [])

        summaries.append(
            ContractSummaryItem(
                contract_id=str(data.get("contract_id") or ""),
                company_id=str(data.get("company_id") or "unknown"),
                timestamp=str(data.get("timestamp") or ""),
                risk_score=int(risk_assessment.get("risk_score", 0) or 0),
                risk_level=str(risk_assessment.get("risk_level") or "medium").lower(),
                violations=residual_violations,
                negotiation_status=str(metrics_service._derive_status(data) or "reviewed"),
            )
        )

    summaries.sort(key=lambda item: item.timestamp, reverse=True)
    return summaries

@router.get("/contracts/deleted", response_model=List[DeletedContractItem])
async def get_deleted_contracts(limit: int = 5):
    deleted = memory.list_deleted_analyses(limit=limit)
    items: List[DeletedContractItem] = []

    for data in deleted:
        risk_assessment = data.get("risk_assessment") or {}
        items.append(
            DeletedContractItem(
                contract_id=str(data.get("contract_id") or ""),
                company_id=str(data.get("company_id") or "unknown"),
                timestamp=str(data.get("timestamp") or ""),
                deleted_at=str(data.get("deleted_at") or ""),
                risk_score=int(risk_assessment.get("risk_score", 0) or 0),
                risk_level=str(risk_assessment.get("risk_level") or "medium").lower(),
                negotiation_status=str(metrics_service._derive_status(data) or "reviewed"),
            )
        )

    return items


@router.get("/contracts/{contract_id}", response_model=ReviewResponse)
async def get_contract_analysis(contract_id: str):
    data = memory.get_analysis(contract_id)
    if not data:
        raise HTTPException(status_code=404, detail="Analysis not found")
    return _ensure_contract_redlines_have_suggestions(data)


@router.delete("/contracts/{contract_id}", response_model=ContractDeleteResponse)
async def delete_contract(contract_id: str, permanent: bool = False):
    deleted = memory.delete_analysis(contract_id, permanent=permanent)
    if not deleted:
        raise HTTPException(status_code=404, detail="Analysis not found")

    return ContractDeleteResponse(deleted=True, contract_id=contract_id, permanent=permanent)


@router.post("/contracts/{contract_id}/restore", response_model=ContractRestoreResponse)
async def restore_contract(contract_id: str):
    restored = memory.restore_analysis(contract_id)
    if not restored:
        raise HTTPException(status_code=404, detail="Analysis not found")
    return ContractRestoreResponse(restored=True, contract_id=contract_id)

@router.get("/contracts/{contract_id}/summary")
async def get_contract_summary(contract_id: str):
    data = memory.get_analysis(contract_id)
    if not data:
        raise HTTPException(status_code=404, detail="Analysis not found")
    return {"summary": data.get("summary")}

@router.get("/contracts/{contract_id}/risks")
async def get_contract_risks(contract_id: str):
    data = memory.get_analysis(contract_id)
    if not data:
        raise HTTPException(status_code=404, detail="Analysis not found")
    return {
        "risk_assessment": data.get("risk_assessment"),
        "violations": data.get("violations"),
        "ai_observations": data.get("ai_observations")
    }

@router.get("/contracts/{contract_id}/redlines")
async def get_contract_redlines(contract_id: str):
    data = memory.get_analysis(contract_id)
    if not data:
        raise HTTPException(status_code=404, detail="Analysis not found")
    data = _ensure_contract_redlines_have_suggestions(data)
    return {"redlines": data.get("redlines")}


@router.put(
    "/contracts/{contract_id}/negotiation-status",
    response_model=NegotiationStatusUpdateResponse,
)
async def update_negotiation_status(contract_id: str, payload: NegotiationStatusUpdateRequest):
    data = memory.get_analysis(contract_id)
    if not data:
        raise HTTPException(status_code=404, detail="Analysis not found")

    allowed = {"pending", "reviewed", "negotiating", "approved"}
    next_status = str(payload.negotiation_status or "").strip().lower()
    if next_status not in allowed:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid negotiation_status '{next_status}'. Allowed: {', '.join(sorted(allowed))}",
        )

    now_iso = datetime.datetime.utcnow().isoformat()
    updates = {
        "negotiation_status": next_status,
        "timestamp": now_iso,
    }

    if next_status == "pending":
        updates["requires_manual_review"] = True
    elif next_status in {"reviewed", "negotiating", "approved"}:
        updates["requires_manual_review"] = False

    if next_status == "negotiating" and not data.get("negotiation_queue_at"):
        updates["negotiation_queue_at"] = now_iso
    if next_status == "pending":
        updates["negotiation_queue_at"] = None

    memory.update_analysis(contract_id, updates)

    return NegotiationStatusUpdateResponse(
        contract_id=contract_id,
        negotiation_status=next_status,
        updated_at=now_iso,
    )


@router.get("/contracts/{contract_id}/download")
async def download_revised_contract(contract_id: str, format: str = "docx"):
    data = memory.get_analysis(contract_id)
    if not data:
        raise HTTPException(status_code=404, detail="Analysis not found")

    data = _ensure_contract_redlines_have_suggestions(data)

    accepted = set(data.get("accepted_redline_indexes", []))
    revised_document = str(data.get("revised_document") or "").strip()
    if not revised_document:
        if accepted:
            revised_document = _rebuild_revised_document(data, accepted).strip()
        else:
            revised_document = str(data.get("document_text") or "").strip()

    if not revised_document:
        raise HTTPException(status_code=400, detail="No document content available for download")

    safe_contract_id = re.sub(r"[^a-zA-Z0-9_-]+", "_", contract_id)
    normalized_format = (format or "docx").strip().lower()

    if normalized_format == "txt":
        filename = f"{safe_contract_id}_revised.txt"
        return Response(
            content=revised_document,
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    if normalized_format != "docx":
        raise HTTPException(status_code=400, detail="Unsupported download format. Use docx or txt")

    try:
        docx_bytes = _build_docx_bytes(revised_document, safe_contract_id)
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx is not installed on backend")
    except Exception as e:
        logger.error(f"Failed generating DOCX for {contract_id}: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate DOCX")

    filename = f"{safe_contract_id}_revised.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/contracts/{contract_id}/redlines/{redline_index}/accept", response_model=RedlineAcceptResponse)
async def accept_redline(contract_id: str, redline_index: int, payload: RedlineAcceptRequest):
    data = memory.get_analysis(contract_id)
    if not data:
        raise HTTPException(status_code=404, detail="Analysis not found")

    redlines = data.get("redlines", [])
    if redline_index < 0 or redline_index >= len(redlines):
        raise HTTPException(status_code=400, detail="Invalid redline index")

    accepted = set(data.get("accepted_redline_indexes", []))
    rejected = set(data.get("rejected_redline_indexes", []))

    decision = (payload.decision or "").strip().lower()
    if decision == "accept":
        accepted.add(redline_index)
        rejected.discard(redline_index)
    elif decision == "reject":
        rejected.add(redline_index)
        accepted.discard(redline_index)
    elif decision == "clear":
        accepted.discard(redline_index)
        rejected.discard(redline_index)
    else:
        # Backward compatibility: accepted=True means accept, accepted=False means clear.
        if payload.accepted:
            accepted.add(redline_index)
            rejected.discard(redline_index)
        else:
            accepted.discard(redline_index)
            rejected.discard(redline_index)

    revised_document = _rebuild_revised_document(data, accepted)
    updated_risk_assessment = _recalculate_risk_assessment_after_redlines(data, accepted)

    all_redlines_reviewed = len(redlines) > 0 and len(accepted.union(rejected)) == len(redlines)
    risk_cleared = int(updated_risk_assessment.get("risk_score", 0) or 0) == 0
    negotiation_status = "negotiating" if (all_redlines_reviewed or risk_cleared) else "reviewed"
    queue_at = data.get("negotiation_queue_at")
    if all_redlines_reviewed and not queue_at:
        queue_at = datetime.datetime.utcnow().isoformat()

    updates = {
        "accepted_redline_indexes": sorted(list(accepted)),
        "rejected_redline_indexes": sorted(list(rejected)),
        "revised_document": revised_document,
        "risk_assessment": updated_risk_assessment,
        "negotiation_status": negotiation_status,
        "negotiation_queue_at": queue_at,
        "requires_manual_review": bool(updated_risk_assessment.get("risk_score", 0) > 60),
    }
    memory.update_analysis(contract_id, updates)

    return RedlineAcceptResponse(
        contract_id=contract_id,
        redline_index=redline_index,
        accepted=redline_index in accepted,
        accepted_redline_indexes=updates["accepted_redline_indexes"],
        rejected_redline_indexes=updates["rejected_redline_indexes"],
        negotiation_status=negotiation_status,
        revised_document=revised_document,
        risk_assessment=updated_risk_assessment,
    )


@router.put(
    "/contracts/{contract_id}/redlines/{redline_index}/suggestion",
    response_model=RedlineSuggestionUpdateResponse,
)
async def update_redline_suggestion(
    contract_id: str,
    redline_index: int,
    payload: RedlineSuggestionUpdateRequest,
):
    data = memory.get_analysis(contract_id)
    if not data:
        raise HTTPException(status_code=404, detail="Analysis not found")

    redlines = data.get("redlines", [])
    if redline_index < 0 or redline_index >= len(redlines):
        raise HTTPException(status_code=400, detail="Invalid redline index")

    suggestion = _sanitize_redline_suggestion(payload.suggested_redline or "")
    if _is_placeholder_suggestion(suggestion):
        raise HTTPException(status_code=400, detail="Suggested redline cannot be empty")

    redlines[redline_index]["suggested_redline"] = suggestion

    accepted = set(data.get("accepted_redline_indexes", []))
    revised_document = _rebuild_revised_document(data, accepted)

    memory.update_analysis(
        contract_id,
        {
            "redlines": redlines,
            "revised_document": revised_document,
        },
    )

    return RedlineSuggestionUpdateResponse(
        contract_id=contract_id,
        redline_index=redline_index,
        suggested_redline=suggestion,
        revised_document=revised_document,
    )


def _rebuild_revised_document(data: dict, accepted_indexes: set[int]) -> str:
    base_document = str(data.get("document_text") or data.get("revised_document") or "")
    if not base_document:
        return ""

    redlines = data.get("redlines", [])
    violations = data.get("violations", [])
    violations_by_clause: dict[str, str] = {}
    for violation in violations:
        clause_key = str(violation.get("clause") or "").strip().lower()
        original_clause = str(violation.get("original_clause") or "").strip()
        if clause_key and original_clause and clause_key not in violations_by_clause:
            violations_by_clause[clause_key] = original_clause

    revised = base_document
    unresolved_suggestions: list[dict[str, str]] = []

    for idx in sorted(accepted_indexes):
        if idx < 0 or idx >= len(redlines):
            continue
        item = redlines[idx]
        original = str(item.get("original_clause") or "").strip()
        suggested = _sanitize_redline_suggestion(
            str(item.get("suggested_redline") or item.get("suggested_clause") or "")
        )
        if _is_placeholder_suggestion(suggested):
            continue
        if not suggested:
            continue

        # If this suggestion is already present, skip replacement for this redline.
        if suggested and suggested in revised:
            continue

        clause_key = str(item.get("violation") or "").strip().lower()
        candidates = _ordered_unique_candidates(
            [
                original,
                violations_by_clause.get(clause_key, ""),
            ]
        )

        changed_for_item = False
        for candidate in candidates:
            updated, changed = _replace_clause_once(revised, candidate, suggested)
            if changed:
                revised = updated
                changed_for_item = True
                break

        if not changed_for_item:
            unresolved_suggestions.append(
                {
                    "clause": str(item.get("violation") or "General"),
                    "original": original or violations_by_clause.get(clause_key, ""),
                    "suggested": suggested,
                }
            )

    if unresolved_suggestions:
        revised = _append_unresolved_revision_notes(revised, unresolved_suggestions)

    return revised


def _ordered_unique_candidates(items: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for item in items:
        value = str(item or "").strip()
        if not value:
            continue
        key = value.lower()
        if key in seen:
            continue
        seen.add(key)
        result.append(value)
    return result


def _replace_clause_once(document: str, original: str, suggested: str) -> tuple[str, bool]:
    if not original or not suggested:
        return document, False

    # Fast path: exact substring replacement.
    if original in document:
        return document.replace(original, suggested, 1), True

    # Case-insensitive literal replacement.
    literal_match = re.search(re.escape(original), document, flags=re.IGNORECASE)
    if literal_match:
        return document[: literal_match.start()] + suggested + document[literal_match.end() :], True

    # Fallback: tolerate whitespace/newline differences between parser outputs.
    flexible_pattern = _to_flexible_whitespace_pattern(original)
    if flexible_pattern:
        match = re.search(flexible_pattern, document, flags=re.IGNORECASE | re.MULTILINE)
        if match:
            return document[: match.start()] + suggested + document[match.end() :], True

    # Fuzzy fallback: replace the most similar paragraph block.
    fuzzy_updated, fuzzy_changed = _replace_fuzzy_paragraph(document, original, suggested)
    if fuzzy_changed:
        return fuzzy_updated, True

    return document, False


def _to_flexible_whitespace_pattern(text: str) -> str:
    tokens = [re.escape(token) for token in re.split(r"\s+", text.strip()) if token]
    if not tokens:
        return ""
    return r"\s+".join(tokens)


def _replace_fuzzy_paragraph(document: str, original: str, suggested: str) -> tuple[str, bool]:
    norm_original = _normalize_for_similarity(original)
    if len(norm_original) < 20:
        return document, False

    best_score = 0.0
    best_span: tuple[int, int] | None = None

    # Candidate blocks are paragraph-like chunks in the source document.
    for match in re.finditer(r"\S(?:.|\n)*?(?=\n\s*\n|\Z)", document):
        block_text = match.group(0)
        norm_block = _normalize_for_similarity(block_text)
        if len(norm_block) < 20:
            continue

        ratio = SequenceMatcher(None, norm_original, norm_block).ratio()
        if ratio > best_score:
            best_score = ratio
            best_span = (match.start(), match.end())

    # Conservative threshold to avoid unrelated replacements.
    if best_span and best_score >= 0.78:
        start, end = best_span
        return document[:start] + suggested + document[end:], True

    return document, False


def _normalize_for_similarity(text: str) -> str:
    normalized = str(text or "")
    normalized = normalized.replace("\u2018", "'").replace("\u2019", "'")
    normalized = normalized.replace("\u201c", '"').replace("\u201d", '"')
    normalized = re.sub(r"\s+", " ", normalized).strip().lower()
    return normalized


def _append_unresolved_revision_notes(document: str, unresolved: list[dict[str, str]]) -> str:
    lines = ["", "", "[Accepted Revisions Pending Exact Placement]"]
    for idx, item in enumerate(unresolved, start=1):
        clause = str(item.get("clause") or "General").strip() or "General"
        original = str(item.get("original") or "").strip()
        suggested = str(item.get("suggested") or "").strip()
        lines.append(f"{idx}. Clause: {clause}")
        if original:
            lines.append(f"   - Original: {original}")
        lines.append(f"   + Suggested: {suggested}")

    suffix = "\n".join(lines)
    if "[Accepted Revisions Pending Exact Placement]" in document:
        # Replace existing unresolved block if present to avoid duplication.
        return re.sub(
            r"\n\n\[Accepted Revisions Pending Exact Placement\][\s\S]*$",
            suffix,
            document,
            flags=re.MULTILINE,
        )
    return document + suffix


def _build_docx_bytes(text: str, contract_id: str) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(f"Revised Contract - {contract_id}", level=1)

    for block in re.split(r"\n\s*\n", text.strip()):
        paragraph = block.strip()
        if paragraph:
            doc.add_paragraph(paragraph)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _recalculate_risk_assessment_after_redlines(data: dict, accepted_indexes: set[int]) -> dict:
    base_risk = data.get("risk_assessment") or {}
    violations = deduplicate_violations(data.get("violations") or [])
    redlines = data.get("redlines") or []

    if not violations:
        return {
            "risk_score": 0,
            "risk_level": "low",
            "violation_counts": {"high": 0, "medium": 0, "low": 0},
        }

    if redlines and len(accepted_indexes) >= len(redlines):
        return {
            "risk_score": 0,
            "risk_level": "low",
            "violation_counts": {"high": 0, "medium": 0, "low": 0},
        }

    fixed_clauses: set[str] = set()
    for idx in accepted_indexes:
        if idx < 0 or idx >= len(redlines):
            continue
        clause = str(redlines[idx].get("violation") or "").strip().lower()
        if clause:
            fixed_clauses.add(clause)

    remaining = []
    for violation in violations:
        clause = str(violation.get("clause") or "").strip().lower()
        if clause and clause in fixed_clauses:
            continue
        remaining.append(violation)

    counts = {"high": 0, "medium": 0, "low": 0}
    for violation in remaining:
        sev = str(violation.get("severity") or "medium").strip().lower()
        if sev not in counts:
            sev = "medium"
        counts[sev] += 1

    if not remaining:
        score = 0
    else:
        # Weighted residual score from remaining violation severities.
        weighted = counts["high"] * 30 + counts["medium"] * 15 + counts["low"] * 5
        score = min(100, max(0, weighted))

        # Smooth drop relative to previous score so accepted fixes visibly reduce risk.
        prev_score = int(base_risk.get("risk_score", score) or score)
        score = min(score, prev_score)

    if score <= 30:
        level = "low"
    elif score <= 60:
        level = "medium"
    else:
        level = "high"

    return {
        "risk_score": int(score),
        "risk_level": level,
        "violation_counts": counts,
    }


@router.post("/chat/docs/upload", response_model=ChatIndexResponse)
async def upload_chat_document(
    file: UploadFile = File(...),
    contract_id: str = "",
    ctx: SecurityContext = Depends(get_security_context),
    _limit: None = Depends(rate_limit_chat),
):
    company_id = ctx.company_id
    rag_service = _get_rag_service()
    if rag_service is None:
        detail = f"RAG service unavailable. {ragsvc_error or ''}".strip()
        raise HTTPException(status_code=503, detail=detail)

    uploaded_path, original_filename = _save_validated_upload(file, prefix_with_uuid=True)

    try:
        result = rag_service.index_document(
            file_path=uploaded_path,
            company_id=company_id,
            contract_id=contract_id or None,
            filename=original_filename,
        )
        return result
    except Exception as e:
        logger.exception("Chat document indexing failed")
        raise HTTPException(status_code=500, detail="Internal server error")


@router.post("/chat/query", response_model=ChatQueryResponse)
async def query_chat_documents(
    payload: ChatQueryRequest,
    ctx: SecurityContext = Depends(get_security_context),
    _limit: None = Depends(rate_limit_chat),
):
    rag_service = _get_rag_service()
    if rag_service is None:
        detail = f"RAG service unavailable. {ragsvc_error or ''}".strip()
        raise HTTPException(status_code=503, detail=detail)

    payload.company_id = ctx.company_id

    try:
        deterministic_company_answer = _try_answer_company_metrics_question(payload)
        if deterministic_company_answer:
            return deterministic_company_answer

        live_context = None
        if payload.contract_id:
            analysis = memory.get_analysis(payload.contract_id)
            if analysis:
                live_context = _build_analysis_context_for_chat(analysis)
        else:
            live_context = _build_company_context_for_chat(payload.company_id)

        result = rag_service.answer_question(
            question=payload.question,
            company_id=payload.company_id,
            contract_id=payload.contract_id,
            doc_id=payload.doc_id,
            top_k=payload.top_k,
            chat_history=[m.model_dump() for m in payload.chat_history],
            additional_context=live_context,
        )

        # If vector retrieval has no hits, fall back to saved contract analysis context.
        if not result.get("grounded", False):
            fallback_result = _answer_from_saved_analysis(rag_service, payload)
            if fallback_result:
                return fallback_result

            if live_context:
                chain = rag_service.prompt | rag_service.llm
                response = chain.invoke(
                    {
                        "question": payload.question,
                        "chat_history": "\n".join(
                            f"{m.role}: {m.content}" for m in (payload.chat_history or [])[-8:]
                        )
                        or "(none)",
                        "context": live_context,
                    }
                )
                return {
                    "question": payload.question,
                    "answer": str(response.content or "").strip(),
                    "citations": [
                        {
                            "chunk_id": f"company:{payload.company_id}",
                            "doc_id": None,
                            "filename": "company_portfolio_context",
                            "contract_id": None,
                            "score": None,
                            "excerpt": live_context[:400],
                        }
                    ],
                    "grounded": False,
                }

        return result
    except Exception as e:
        logger.exception("Chat query failed")
        raise HTTPException(status_code=500, detail="Internal server error")


@router.get("/metrics/dashboard", response_model=DashboardMetricsResponse)
async def get_dashboard_metrics(timeframe: str = "all"):
    try:
        return metrics_service.get_dashboard_metrics(timeframe=timeframe)
    except Exception as e:
        logger.exception("Dashboard metrics failed")
        raise HTTPException(status_code=500, detail="Internal server error")


@router.get("/metrics/analytics", response_model=AnalyticsMetricsResponse)
async def get_analytics_metrics(timeframe: str = "all"):
    try:
        return metrics_service.get_analytics_metrics(timeframe=timeframe)
    except Exception as e:
        logger.exception("Analytics metrics failed")
        raise HTTPException(status_code=500, detail="Internal server error")


@router.get("/metrics/negotiations", response_model=List[NegotiationItem])
async def get_negotiations_metrics(timeframe: str = "all"):
    try:
        return metrics_service.get_negotiations(timeframe=timeframe)
    except Exception as e:
        logger.exception("Negotiations metrics failed")
        raise HTTPException(status_code=500, detail="Internal server error")
